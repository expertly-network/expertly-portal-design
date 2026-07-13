from docx import Document

doc = Document()
doc.add_heading("Expertly: Hero Headlines", level=0)

intro = doc.add_paragraph()
intro.add_run(
    "Final hero copy. Lines rotate in sequence: common, then practitioner, then client, then "
    "back to common, so the loop repeats indefinitely. Both audience lines share the same "
    "quiet-confidence register: a short affirmation, then a payoff, with neither line sounding "
    "like marketing copy aimed at 'a client' or 'a practitioner', each just sounds like it's "
    "talking to a person."
).italic = True
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("1. Common: ").bold = True
p.add_run("You Bring the Expertise. They Bring the Need. We Bring You Together.")
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("2. Practitioner: ").bold = True
p.add_run("You're Good at This. Let the Right People Find Out.")
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("3. Client: ").bold = True
p.add_run("You're Stuck on Something. Someone Here Already Solved It.")
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("4. Loops back to Common, repeats.").italic = True
doc.add_paragraph("")

doc.add_heading("Hero: Supporting Copy", level=1)

p = doc.add_paragraph()
p.add_run("Subhead: ").bold = True
p.add_run(
    "Every expert here was checked before they got listed. Every search here is built to find "
    "exactly who you need: no broker fees, no cold outreach, no guessing who's actually good "
    "at this."
)
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("Search placeholder: ").bold = True
p.add_run("Cross-border M&A counsel in Dubai, fluent in English & Arabic")
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("Hint chips (\"Try\"): ").bold = True
p.add_run("Transfer pricing · Singapore  /  Antitrust counsel · UK  /  Compliance lead · UAE")
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("Trust stats: ").bold = True
doc.add_paragraph("20+: Already listed, already checked")
doc.add_paragraph("20+: Countries reached")
doc.add_paragraph("12: Practice areas, not one generalist label")
doc.add_paragraph("100+: Real matches made")
doc.add_paragraph("")

p = doc.add_paragraph()
p.add_run("CTA Buttons: ").bold = True
p.add_run(
    "Always visible (not tied to the rotating headline, so they're never mid-fade when "
    "someone tries to click). Placed as a pair right after the trust stats, closing out "
    "the hero. Sequence is: subhead → AI search bar → trust stats → these two buttons."
)
doc.add_paragraph("")
doc.add_paragraph("Client (primary / filled button): \"Find an Expert\"")
doc.add_paragraph("Practitioner (secondary / outline button): \"Become One\"")
p = doc.add_paragraph()
p.add_run(
    "The pairing answers itself: \"Find an Expert\" / \"Become One\", echoing the main "
    "headline's 'they bring the need, you bring the expertise' idea without restating it."
).italic = True

filename = "Expertly_Hero_Headlines.docx"
doc.save(filename)
print("Saved", filename)
